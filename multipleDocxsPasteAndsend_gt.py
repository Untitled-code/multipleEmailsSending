# a program that multiplies word documents with different addresses and send it by email

from pathlib import Path
from docx import Document
import os, logging
import subprocess
import re
import time
import csv
import smtplib
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email import encoders
from email.mime.text import MIMEText # для работы с кириллицей
from email.header import Header
logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s - %(message)s')
logging.debug('Start of program')
QUERY_LIST_FILE_EMAILS = Path(f'./your_emails.csv')
QUERY_LIST_FILE_REGIONS = Path(f'./your_regions_list.csv') #names of addresses
#body of text message
body = "Доброго дня,\nНадсилаю інформаційний запит згідно ЗУ Про доступ до публічної інформації.\nПрошу зареєструвати та повідомити вхідний номер мого запиту.\nДякую за розуміння та співпрацю!\nЗ повагою,\nОксана Орсач"
def get_keywords(query_file):
    with open(query_file, 'r') as i_file:
        rows = csv.reader(i_file, delimiter=',')
        keywords = [row[0] for row in rows]
        return keywords

def get_para_data(output_doc_name, paragraph): #function that copy whole word doc with styles
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

def sendEmail(sending_file, file_name, email):
    # LOg in to email account
    print("Logging into email account")
    time.sleep(2)
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    # pass updated
    smtpObj.login('yourMail@mail.com', 'password')
    # send out reminder letters
    """ Sending email with file attached """
    msg = MIMEMultipart()
    msg['Subject'] = Header('Запит на доступ до публічної інформації', 'utf-8')
    msg.attach(MIMEText(body.encode('utf-8'), _charset='utf-8'))  #attach text message
    attachment = open(sending_file, 'rb')
    part = MIMEBase("application", "octet-stream")
    part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f"attachment; filename= {file_name}") #attaching a file
    msg.attach(part)

    print('Sending email to %s...' % email)
    logging.debug('Sending email to %s...' % email)
    sendmailStatus = smtpObj.sendmail('your', email, msg.as_string())
    if sendmailStatus != {}:
        print('There was problem sending email to %s: %s' % (email, sendmailStatus))
    smtpObj.quit()

def formingDocx(i):
    input_doc = Document('./zapyt.docx') #your template of request
    output_doc = Document()
    stringToPut = f'До ГУНП в {REGIONS[i]} області' # changing the regions in the requests
    paragraph = output_doc.add_paragraph(stringToPut) # pasting the string to the doc file
    paragraph.alignment = 2 #alignment text to the right
    # Call the function of copying the doc
    for para in input_doc.paragraphs:
        get_para_data(output_doc, para)
    output_name = f'zapytGUNP{str(i)}.docx'
    output_name_fullPath = f'./{output_name}'
    output_doc.save(output_name_fullPath)
    print(f"New word file saved...{output_name}")
    sending_file = output_name_fullPath
    sendEmail(sending_file, output_name, EMAILS[i])


###################
REGIONS = get_keywords(QUERY_LIST_FILE_REGIONS)
EMAILS = get_keywords(QUERY_LIST_FILE_EMAILS)
for i in range(len(EMAILS)):
    formingDocx(i)
print("Done!")